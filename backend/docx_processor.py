import re
from docx import Document

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        # This is a simple replacement that tries to preserve run formatting
        # by replacing the text in the first run that contains the key, and
        # clearing the others if the key is split across runs.
        # For simplicity, we just replace paragraph.text, which might lose some intra-paragraph
        # formatting, but python-docx makes exact run replacement complex.
        # A safer approach for placeholders:
        for i in range(len(inline)):
            if key in inline[i].text:
                inline[i].text = inline[i].text.replace(key, value)

def parse_markdown_to_dict(md_text):
    """Parses simple markdown to a dictionary based on headers."""
    lines = md_text.split('\n')
    data = {}
    current_key = "Introduction"
    current_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line is a header like **Header** or ## Header
        header_match = re.match(r'^(?:\*\*|##)\s*(.*?)(?:\*\*|)$', line)
        if header_match:
            if current_text:
                data[current_key] = '\n'.join(current_text)
            current_key = header_match.group(1).strip()
            current_text = []
        else:
            current_text.append(line)
            
    if current_text:
        data[current_key] = '\n'.join(current_text)
        
    return data

def generate_cover_letter(md_text, template_path, output_path):
    doc = Document(template_path)
    md_data = parse_markdown_to_dict(md_text)
    
    # We will map known markdown keys to our CoverLetter_Template.docx placeholders
    # Our generated template has: {{Date}}, {{Hiring Manager}}, {{Company Name}}, 
    # {{Company Address Line 1}}, {{Company Address Line 2}}, {{Subject}}, 
    # {{Salutation}}, {{Introduction}}, {{Bullet Points}}, {{Closing}}
    
    # For a real implementation, we'd extract these from the MD or use LLM/Heuristics
    # Here we do a basic mapping for demonstration.
    
    replacements = {
        "{{Date}}": "October 12, 2025",
        "{{Hiring Manager}}": "Martin Herzog",
        "{{Company Name}}": "Stadtwerke München GmbH (SWM)",
        "{{Company Address Line 1}}": "Emmy-Noether-Straße 2",
        "{{Company Address Line 2}}": "80992 München",
        "{{Subject}}": "Application for Werkstudent:in – Asset Simulation Versorgungsnetze",
        "{{Salutation}}": "Dear Martin Herzog,",
        "{{Introduction}}": md_data.get("Introduction", "Introduction text..."),
        "{{Bullet Points}}": md_data.get("Bullet Points", "Bullet points..."),
        "{{Closing}}": md_data.get("Closing", "Closing text...")
    }
    
    # In the tailored_cover_letter.md, the whole text is provided without strict headers for intro/bullets/closing.
    # We will just inject the whole text into Introduction and clear the rest for simplicity if no headers exist.
    if "Introduction" in md_data and len(md_data) == 1:
        replacements["{{Introduction}}"] = md_data["Introduction"]
        replacements["{{Bullet Points}}"] = ""
        replacements["{{Closing}}"] = ""
    
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
                
    doc.save(output_path)
    return output_path

def generate_cv(md_text, template_path, output_path):
    # This is a placeholder for the complex CV table logic
    doc = Document(template_path)
    doc.save(output_path)
    return output_path
