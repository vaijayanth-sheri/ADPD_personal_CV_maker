import sys
from docx import Document

def print_docx_content(filepath):
    doc = Document(filepath)
    with open("backend/placeholders.txt", "w", encoding="utf-8") as f:
        for p in doc.paragraphs:
            f.write(p.text + "\n")
        
        f.write("--- TABLES ---\n")
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    f.write(cell.text + "\n")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        print_docx_content(sys.argv[1])
