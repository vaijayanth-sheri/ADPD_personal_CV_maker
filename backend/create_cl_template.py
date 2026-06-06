from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add Header (Name and contact info)
name_p = doc.add_paragraph()
name_run = name_p.add_run("Vaijayanth Sheri M.Eng")
name_run.font.size = Pt(24)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(0x2E, 0x40, 0x53)
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

contact_p = doc.add_paragraph()
contact_run = contact_p.add_run("sheri.vaijayanth@gmail.com | +49 17648060315 | Semperstr. 10, München")
contact_run.font.size = Pt(10)
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph() # spacing

# Add Date
date_p = doc.add_paragraph("{{Date}}")
date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add Hiring Manager Info
doc.add_paragraph("{{Hiring Manager}}")
doc.add_paragraph("{{Company Name}}")
doc.add_paragraph("{{Company Address Line 1}}")
doc.add_paragraph("{{Company Address Line 2}}")

doc.add_paragraph() # spacing

# Subject
subj_p = doc.add_paragraph()
subj_run = subj_p.add_run("{{Subject}}")
subj_run.font.bold = True

doc.add_paragraph()

# Salutation
doc.add_paragraph("{{Salutation}}")

doc.add_paragraph()

# Introduction
doc.add_paragraph("{{Introduction}}")

doc.add_paragraph()

# Bullet Points
doc.add_paragraph("{{Bullet Points}}")

doc.add_paragraph()

# Closing
doc.add_paragraph("{{Closing}}")

doc.add_paragraph()

# Sign-off
doc.add_paragraph("Thank you for your time and consideration.")
doc.add_paragraph()
doc.add_paragraph("Sincerely,")
doc.add_paragraph("Vaijayanth Sheri")

doc.save("templates/CoverLetter_Template.docx")
print("Cover letter template created successfully.")
